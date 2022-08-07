import os
import json
import re
import requests
import bs4
import docx
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
import time
import requests

chrome_options = Options()
chrome_options.add_experimental_option("detach", True)
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
PROBLEMSET_BASE_URL = "https://leetcode.com/problemset/all/?page="

def save_header(document, problem_name):
    heading = document.add_heading(problem_name,0)
    heading.alignment = 1
    title_style = heading.style
    title_style.font.size = docx.shared.Pt(16)
    title_style.font.bold = True
    title_style.font.name = 'Consolas'
    title_style.font.all_caps = True

def save_to_docx(document,problem_name, description):
    save_header(document, problem_name)
    soup = bs4.BeautifulSoup(description, "html.parser")
    blank =  soup.find_all("p", text=re.compile('\xa0'))
    for i in blank:
        i.decompose()

    for p in soup.find_all("p"):
        if "Example" in p.text or "Note" in p.text or "Constraints" in p.text:
            tx = document.add_paragraph()
            runner = tx.add_run(p.text.strip())
            runner.bold = True
            runner.underline = True
            runner.space_after = False
            runner.space_before = False
            runner.line_spacing_rule = 1

            if "Example" in p.text:
                if p.next_sibling.next_sibling.name == "img":
                    response = requests.get(p.next_sibling.next_sibling["src"]) 
                    binary_img = BytesIO(response.content) 
                    document.add_picture(binary_img, width=docx.shared.Inches(3))
                else:
                    code_text = p.next_sibling.next_sibling
                    document.add_paragraph(code_text.text.strip())
            elif "Constraints" in p.text:
                constraints = p.next_sibling.next_sibling
                for element in constraints.findAll('li'):
                    if element.find('sup'):
                        for sup in element.findAll('sup'):
                            sup.replaceWith('^' + sup.text)
                    document.add_paragraph(element.text.strip(), style='List Bullet')
                    # cst = element.text.split('^')
                    # tx = document.add_paragraph(cst[0], style='List Bullet')
                    # for c in range(1,len(cst)):
                    #     k = 0
                    #     while cst[c][k].isnumeric():
                    #         runner = tx.add_run(cst[c][k])
                    #         runner.font.superscript = True
                    #         k+=1
                    #     afterrunner = tx.add_run(cst[c][1:])
                    #     afterrunner.font.superscript = False
            else:
                code_text = p.next_sibling.next_sibling
                document.add_paragraph(code_text.text.strip())
        else:
            document.add_paragraph(p.text.strip())
            

def get_html(url):
    r = requests.get(url)
    if r.status_code != 200:
        print("Error: Could not get problem page")
        return
    html = r.content
    return html

def get_problem_desctiption(document,url, problem_name, premium):
    if premium:
        save_header(document, problem_name)
        document.add_paragraph("Premium Problem")
        return
    html = driver.get(url)
    
    WebDriverWait(driver, 60).until(EC.invisibility_of_element_located((By.ID, "initial-loading")))
    html = driver.page_source
    soup = bs4.BeautifulSoup(html, "html.parser")
    description = soup.find("div", {"class": "content__u3I1 question-content__JfgR"})
    save_to_docx(document, problem_name, str(description.contents))
    
def main():
    document = docx.Document()
    problemset= []

    if os.path.exists("problemset.docx"):
        document = docx.Document('problemset.docx')
    
    if os.path.exists("problemset.json"):
        with open('problemset.json', 'r') as f:
            problemset = json.load(f)

    for i in range(1,49):
        url = PROBLEMSET_BASE_URL + str(i)
        print("processing page: " + url)
        html = driver.get(url)
        time.sleep(2)
        html = driver.page_source
        soup = bs4.BeautifulSoup(html, "html.parser")
        table = soup.find("div", {"role": "rowgroup"})
        rows = table.find_all("div", {"role": "row"})
        for row in rows:
            cells = row.find_all("div", {"role": "cell"})
            problem = {}
            print(cells[1].text)
            problem['premium'] = False
            problem["title"] = cells[1].text
            problem["url"] = "https://leetcode.com" + cells[1].find("a")["href"]
            problem["Acceptance"] = cells[3].text
            problem["difficulty"] = cells[4].text
            if cells[0].find("svg").find("path").get("d") == "M19 11.063V7h-2v1a1 1 0 11-2 0V7H9v1a1 1 0 01-2 0V7H5v4.063h14zm0 2H5V19h14v-5.938zM9 5h6V4a1 1 0 112 0v1h2a2 2 0 012 2v12a2 2 0 01-2 2H5a2 2 0 01-2-2V7a2 2 0 012-2h2V4a1 1 0 012 0v1z":
                print("Ignore Daily Challenge Problem")
                continue
            elif cells[0].find("svg").find("path").get("d") == "M7 8v2H6a3 3 0 00-3 3v6a3 3 0 003 3h12a3 3 0 003-3v-6a3 3 0 00-3-3h-1V8A5 5 0 007 8zm8 0v2H9V8a3 3 0 116 0zm-3 6a2 2 0 100 4 2 2 0 000-4z":
                problem['premium'] = True
            try:
                get_problem_desctiption(document, problem["url"], problem["title"], problem["premium"])
            except Exception as e:
                print("Error: Could not get problem page: ", problem["url"], problem["title"], problem["premium"])
                i-=1
                continue
            problemset.append(problem)
    
    document.save("problemset.docx")
    with open('problemset.json', 'w') as f:
        json.dump(problemset, f)

main()